"""
Generates one test fixture per remaining heterogeneous format the
manager flagged: a Word document, a standalone photographed image,
and a true scanned (image-only, no text layer) PDF -- built by
rasterizing an existing digital PDF and re-saving it as images only,
so the ingestion agent's "digital vs scanned" detection is genuinely
exercised rather than assumed.

Run: python3 samples/generate_other_formats.py
"""

import io
from pathlib import Path

import docx
import fitz  # PyMuPDF
import img2pdf
import numpy as np
from PIL import Image, ImageDraw, ImageFilter, ImageFont

OUT_DIR = Path("samples/other_formats")


def generate_docx_correspondence():
    """A claimant follow-up letter as a real .docx -- no native page
    coordinates, exercising the bbox=None / paragraph-index fallback
    path in ingest_docx()."""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = docx.Document()
    doc.add_heading("Claim Correspondence", level=1)
    doc.add_paragraph("Claim ID: AUTO-001")
    doc.add_paragraph("Policy Number: AUTO-2026-58213")
    doc.add_paragraph("Claimant Name: Priya Nair")
    doc.add_paragraph("")
    doc.add_paragraph(
        "Dear Ms. Nair, we are writing to inform you that your claim is currently under "
        "review by our adjuster team. We have received your repair estimate and are "
        "cross-checking it against the inspection report submitted earlier this month."
    )
    doc.add_paragraph(
        "Please note that any additional supporting documents, including updated repair "
        "quotes or photographs, should be submitted within 10 business days to avoid "
        "delays in claim settlement."
    )
    doc.add_paragraph("Regards,\nClaims Department")
    path = OUT_DIR / "claimant_followup_letter.docx"
    doc.save(path)
    print(f"Generated {path}")
    return path


def generate_standalone_image():
    """Simulates a phone-photographed document (e.g. a handwritten
    receipt or a damage photo with an overlaid caption) -- a plain
    image with no PDF wrapper at all, with mild blur/rotation to
    mimic a real photo rather than a clean screenshot."""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (900, 500), color=(250, 248, 245))
    draw = ImageDraw.Draw(img)
    try:
        font = ImageFont.truetype("/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf", 24)
    except Exception:
        font = ImageFont.load_default()

    lines = [
        "RECEIPT - AutoFix Garage Pvt. Ltd.",
        "Claim Ref: AUTO-2026-58213",
        "Date: 2026-04-02",
        "",
        "Front bumper replacement ........ USD 8,500",
        "Labor (4 hrs) .................... USD 6,000",
        "",
        "TOTAL PAID: USD 14,500",
    ]
    y = 40
    for line in lines:
        draw.text((40, y), line, fill=(20, 20, 20), font=font)
        y += 40

    # Mild rotation + blur to mimic an actual phone photo rather than a
    # perfect digital render -- this is what makes Tesseract's job
    # nontrivial (and what makes confidence scores meaningful).
    img = img.rotate(1.2, expand=True, fillcolor=(250, 248, 245))
    img = img.filter(ImageFilter.GaussianBlur(radius=0.6))

    path = OUT_DIR / "repair_receipt_photo.png"
    img.save(path)
    print(f"Generated {path}")
    return path


def generate_scanned_pdf(source_digital_pdf: str, out_name: str, dpi: int = 200, noise_level: int = 6):
    """Takes an existing DIGITAL pdf (with a text layer) and produces a
    true SCANNED-style pdf: rasterize every page to an image, add
    light noise/rotation to mimic a real scanner, then reassemble as a
    PDF made entirely of images (zero extractable text). This is what
    proves the is_digital_pdf() check in ocr_agent.py is actually
    discriminating, not just trusting the .pdf extension."""
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = fitz.open(source_digital_pdf)
    image_bytes_list = []

    rng = np.random.default_rng(7)
    zoom = dpi / 72
    matrix = fitz.Matrix(zoom, zoom)

    for page in doc:
        pix = page.get_pixmap(matrix=matrix)
        img = Image.frombytes("RGB", (pix.width, pix.height), pix.samples)

        # Add scanner-like noise + slight skew.
        arr = np.array(img).astype(np.int16)
        noise = rng.integers(-noise_level, noise_level, arr.shape, dtype=np.int16)
        arr = np.clip(arr + noise, 0, 255).astype(np.uint8)
        img = Image.fromarray(arr)
        img = img.rotate(0.4, expand=True, fillcolor=(255, 255, 255))
        img = img.filter(ImageFilter.GaussianBlur(radius=0.4))

        buf = io.BytesIO()
        img.save(buf, format="JPEG", quality=85)
        image_bytes_list.append(buf.getvalue())

    doc.close()

    out_path = OUT_DIR / out_name
    with open(out_path, "wb") as f:
        f.write(img2pdf.convert(image_bytes_list))
    print(f"Generated {out_path} ({len(image_bytes_list)} pages, image-only)")
    return out_path


def main():
    generate_docx_correspondence()
    generate_standalone_image()
    generate_scanned_pdf("samples/corpus/auto/AUTO-002.pdf", "AUTO-002_scanned.pdf")


if __name__ == "__main__":
    main()
