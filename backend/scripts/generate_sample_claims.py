"""
scripts/generate_sample_claims.py
----------------------------------
Builds `samples/claim_folders/` -- nine realistic, synthetic claim
folders (3 per line of business: auto, property, health) covering the
scenarios that matter for testing the pipeline end-to-end:

  {lob}_complete_clean/     every mandatory doc present, fields agree ->
                             should land STP candidate.
  {lob}_incomplete/         mandatory doc types missing -> should land
                             needs_review / high_risk_incomplete, with
                             missing_mandatory_docs populated.
  {lob}_conflicting_values/ two documents disagree on the same dollar
                             amount -> should surface as HIGH_RISK with
                             a CONFLICTING field status.

Every folder also nests its documents under subfolders (fnol/, evidence/,
correspondence/, etc.) to exercise recursive/webkitdirectory folder
upload, and every folder includes at least one genuinely irrelevant file
(wrong topic entirely, or an unsupported extension) to confirm the
pipeline doesn't get confused by noise -- it should either tag it
"unknown" or record an "INGESTION FAILED" warning without derailing
classification of the real documents.

`auto_conflicting_values/` additionally includes the existing 22-page
`samples/long_supplement_22pages.pdf` to cover the long multi-page PDF
case in the same run.

Usage:
    python scripts/generate_sample_claims.py

Idempotent -- re-running deletes and rebuilds samples/claim_folders/ from
scratch, so it's always safe to just run it again after editing this
script.
"""

from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from PIL import Image, ImageDraw, ImageFont

REPO_ROOT = Path(__file__).resolve().parent.parent
SAMPLES = REPO_ROOT / "samples"
OUT_ROOT = SAMPLES / "claim_folders"


# ---------------------------------------------------------------------------
# Small helpers
# ---------------------------------------------------------------------------


def _doc(path: Path, title: str, paragraphs: list[str]) -> None:
    """Writes a simple DOCX: a bold title heading, then one paragraph per
    line of body text. Deliberately plain formatting -- python-docx's
    ingestion (agents/ingestion/docx_extractor.py) reads paragraph text,
    not layout, so fancy styling wouldn't add test value."""
    path.parent.mkdir(parents=True, exist_ok=True)
    d = Document()
    d.add_heading(title, level=1)
    for para in paragraphs:
        d.add_paragraph(para)
    d.save(path)


def _photo(path: Path, caption: str, size: tuple[int, int] = (640, 480)) -> None:
    """A placeholder "photo" -- a solid-color rectangle with a caption
    burned in. Stands in for real damage/scene photos; what matters for
    pipeline testing is that it's a valid, OCR-able image the crop/bbox
    machinery can render, not that it looks like a real photograph."""
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", size, color=(60, 66, 78))
    d = ImageDraw.Draw(img)
    d.rectangle([20, 20, size[0] - 20, size[1] - 20], outline=(217, 178, 76), width=4)
    try:
        font = ImageFont.load_default(size=20)
    except TypeError:
        font = ImageFont.load_default()
    d.multiline_text((40, size[1] // 2 - 20), caption, fill=(240, 240, 240), font=font)
    img.save(path)


def _irrelevant_docx(path: Path) -> None:
    _doc(
        path,
        "Q3 All-Hands Notes",
        [
            "Thanks everyone for joining the all-hands. Kitchen renovation on the "
            "3rd floor starts Monday -- please use the 2nd floor kitchen until further notice.",
            "Open enrollment for benefits closes on the 15th. Reach out to HR with questions.",
            "Reminder: badge access to the parking garage now requires the new app.",
        ],
    )


def _irrelevant_unsupported(path: Path) -> None:
    """An unsupported extension (.csv isn't in SUPPORTED_FILE_EXTENSIONS)
    -- exercises the "ingestion fails gracefully" path rather than the
    "wrong topic but readable" path covered by _irrelevant_docx."""
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text("item,quantity,unit_cost\nstaples,4,3.29\ncoffee,12,8.50\n")


# ---------------------------------------------------------------------------
# AUTO
# ---------------------------------------------------------------------------


def build_auto_complete_clean(root: Path) -> None:
    _doc(
        root / "fnol" / "notice_of_loss.docx",
        "Notice of Loss - Automobile",
        [
            "Policy Number: AUTO-7734-9021",
            "Carrier Name: Meridian Mutual Insurance",
            "Policy Effective Dates: 01/01/2026 to 01/01/2027",
            "Date of Loss: 03/14/2026",
            "Time of Loss: approximately 5:45 PM",
            "Loss Location: Intersection of 5th Ave and Cedar St, Springfield, IL",
            "Loss Description: Insured vehicle was struck on the rear passenger side "
            "while stopped at a red light by another vehicle that failed to brake in time.",
            "Vehicle VIN: 1HGCM82633A004352",
            "Vehicle Make/Model/Year: 2023 Honda Accord",
            "Damage Location: Rear passenger quarter panel and bumper",
            "Driver Name: Maria Gonzalez",
            "Driver License Number: IL-G4471-0293",
            "Other Party Name: Thomas Reid",
            "Other Vehicle Plate: IL-8829-JR",
            "Witness Name: Patrick Uyeda",
            "Injuries Reported: None reported at the scene.",
            "Rental Needed: Yes, insured requested a rental vehicle during repairs.",
        ],
    )
    _doc(
        root / "police_report" / "police_report.docx",
        "Springfield Police Department - Incident Report",
        [
            "Police Report Number: SPD-2026-011824",
            "Reporting Department: Springfield Police Department, Traffic Division",
            "Date of Loss: 03/14/2026, approximately 5:45 PM",
            "Location: 5th Ave and Cedar St, Springfield, IL",
            "Narrative: Vehicle 2 (Reid) failed to stop for traffic signal and struck "
            "the rear of Vehicle 1 (Gonzalez), which was stationary. No injuries reported "
            "by either party at the scene.",
        ],
    )
    _doc(
        root / "repair_estimate" / "repair_estimate.docx",
        "Springfield Auto Body - Repair Estimate",
        [
            "Estimate prepared for: Maria Gonzalez, 2023 Honda Accord",
            "Repair Estimate Amount: $4,250.00",
            "Scope: Replace rear bumper cover, repair quarter panel, blend paint.",
        ],
    )
    _doc(
        root / "policy_declaration" / "policy_declaration.docx",
        "Meridian Mutual Insurance - Policy Declarations",
        [
            "Policy Number: AUTO-7734-9021",
            "Named Insured: Maria Gonzalez",
            "Policy Effective Dates: 01/01/2026 to 01/01/2027",
            "Coverage: Comprehensive, Collision, Liability",
        ],
    )
    _photo(root / "evidence" / "rear_bumper_damage.jpg", "Rear bumper damage - Gonzalez claim")
    _photo(root / "evidence" / "scene_wide_shot.jpg", "Intersection scene, 5th Ave & Cedar St")
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


def build_auto_incomplete(root: Path) -> None:
    # Missing mandatory: police_report, policy_declaration.
    _doc(
        root / "fnol" / "notice_of_loss.docx",
        "Notice of Loss - Automobile",
        [
            "Policy Number: AUTO-5512-3387",
            "Carrier Name: Meridian Mutual Insurance",
            "Date of Loss: 05/02/2026",
            "Loss Location: Parking garage, 400 Oak St, Denver, CO",
            "Loss Description: Insured's vehicle was sideswiped by an unknown vehicle "
            "while parked; no witnesses or other party information available.",
            "Vehicle VIN: 5YJSA1E26MF123456",
            "Vehicle Make/Model/Year: 2021 Tesla Model S",
            "Damage Location: Driver side doors",
            "Driver Name: Andre Whitfield",
            "Injuries Reported: None.",
        ],
    )
    _doc(
        root / "repair_estimate" / "repair_estimate.docx",
        "Denver Collision Center - Repair Estimate",
        [
            "Estimate prepared for: Andre Whitfield, 2021 Tesla Model S",
            "Repair Estimate Amount: $6,180.00",
            "Scope: Repaint and repair driver side doors, replace side mirror.",
        ],
    )
    _photo(root / "evidence" / "driver_side_scrape.jpg", "Driver side scrape - Whitfield claim")
    _irrelevant_unsupported(root / "irrelevant" / "unrelated_expense_log.csv")


def build_auto_conflicting_values(root: Path) -> None:
    _doc(
        root / "fnol" / "notice_of_loss.docx",
        "Notice of Loss - Automobile",
        [
            "Policy Number: AUTO-9902-1145",
            "Carrier Name: Meridian Mutual Insurance",
            "Policy Effective Dates: 06/01/2025 to 06/01/2026",
            "Date of Loss: 02/09/2026",
            "Time of Loss: approximately 8:15 AM",
            "Loss Location: Highway 40 eastbound, mile marker 112",
            "Loss Description: Multi-vehicle collision in heavy fog; insured vehicle "
            "sustained front-end damage after rear-ending another vehicle that braked suddenly.",
            "Vehicle VIN: 2T1BURHE0JC014589",
            "Vehicle Make/Model/Year: 2018 Toyota Corolla",
            "Damage Location: Front bumper, hood, headlight assembly",
            "Driver Name: Priya Chandrasekaran",
            "Driver License Number: CO-C8827-1104",
            "Other Party Name: Leon Marsh",
            "Injuries Reported: Minor whiplash reported by insured driver.",
            "Injury Description: Insured complained of neck stiffness; declined "
            "on-scene medical transport, sought urgent care same day.",
        ],
    )
    _doc(
        root / "police_report" / "police_report.docx",
        "Colorado State Patrol - Crash Report",
        [
            "Police Report Number: CSP-2026-004471",
            "Reporting Department: Colorado State Patrol",
            "Date of Loss: 02/09/2026, approximately 8:15 AM",
            "Location: Highway 40 eastbound, mile marker 112",
            "Narrative: Reduced visibility due to fog contributed to a three-vehicle "
            "chain collision. Vehicle 2 (Chandrasekaran) rear-ended Vehicle 1 after "
            "Vehicle 1 braked for stopped traffic ahead.",
        ],
    )
    _doc(
        root / "repair_estimate" / "estimate_riverside_auto.docx",
        "Riverside Auto Repair - Estimate",
        [
            "Estimate prepared for: Priya Chandrasekaran, 2018 Toyota Corolla",
            "Repair Estimate Amount: $3,420.00",
            "Scope: Front bumper replacement, hood repair, headlight assembly replacement.",
        ],
    )
    _doc(
        root / "repair_estimate" / "estimate_mountainview_collision.docx",
        "Mountain View Collision - Estimate",
        [
            "Estimate prepared for: Priya Chandrasekaran, 2018 Toyota Corolla",
            "Repair Estimate Amount: $5,150.00",
            "Scope: Front bumper and hood replacement, headlight assembly, radiator "
            "support inspection, frame alignment check.",
        ],
    )
    _doc(
        root / "policy_declaration" / "policy_declaration.docx",
        "Meridian Mutual Insurance - Policy Declarations",
        [
            "Policy Number: AUTO-9902-1145",
            "Named Insured: Priya Chandrasekaran",
            "Policy Effective Dates: 06/01/2025 to 06/01/2026",
            "Coverage: Comprehensive, Collision, Liability, Medical Payments",
        ],
    )
    _photo(root / "evidence" / "front_end_damage.jpg", "Front-end damage - Chandrasekaran claim")
    long_supplement = SAMPLES / "long_supplement_22pages.pdf"
    if long_supplement.exists():
        dest = root / "supplement" / "adjuster_file_notes_22pages.pdf"
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(long_supplement, dest)
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


# ---------------------------------------------------------------------------
# PROPERTY
# ---------------------------------------------------------------------------


def build_property_complete_clean(root: Path) -> None:
    _doc(
        root / "fnol" / "notice_of_loss.docx",
        "Notice of Loss - Property",
        [
            "Policy Number: PROP-4471-0092",
            "Carrier Name: Harborstone Insurance Group",
            "Policy Effective Dates: 03/01/2026 to 03/01/2027",
            "Date of Loss: 04/22/2026",
            "Cause of Loss: Kitchen fire originating from an unattended stovetop.",
            "Loss Description: Fire caused significant smoke and heat damage to the "
            "kitchen and adjoining dining room; sprinkler activation caused water "
            "damage to the hallway flooring.",
            "Property Address: 118 Birchwood Lane, Fairview, OR",
            "Property Type: Single-family residence",
            "Occupancy Status: Owner-occupied, primary residence",
        ],
    )
    _doc(
        root / "inspection_report" / "inspection_report.docx",
        "Harborstone Field Inspection Report",
        [
            "Inspection Report Reference: HFI-2026-3390",
            "Fire Department Report Number: FVFD-2026-00871",
            "Findings: Fire confined to kitchen and dining room; structural framing "
            "intact. Smoke damage extends to hallway and living room ceiling. Water "
            "damage to hallway subfloor from sprinkler discharge.",
        ],
    )
    _doc(
        root / "inventory_list" / "inventory_list.docx",
        "Damaged Contents Inventory",
        [
            "Inventory Reference: INV-2026-3390-A",
            "Damaged Items List: Kitchen cabinets (upper and lower), range/oven, "
            "refrigerator, dining table and 6 chairs, hallway runner rug, ceiling drywall "
            "in kitchen/dining/hallway.",
        ],
    )
    _doc(
        root / "receipts" / "receipts.docx",
        "Loss Valuation Summary",
        [
            "Probable Amount of Loss: $38,600.00",
            "Replacement Cost Estimate: $41,200.00",
            "Actual Cash Value Estimate: $35,750.00",
            "Deductible Amount: $1,000.00",
        ],
    )
    _doc(
        root / "policy_declaration" / "policy_declaration.docx",
        "Harborstone Insurance Group - Policy Declarations",
        [
            "Policy Number: PROP-4471-0092",
            "Named Insured: Daniel Okafor",
            "Policy Effective Dates: 03/01/2026 to 03/01/2027",
            "Coverage: Dwelling, Personal Property, Loss of Use",
        ],
    )
    _photo(root / "evidence" / "kitchen_fire_damage.jpg", "Kitchen fire damage - Okafor claim")
    _photo(root / "evidence" / "hallway_water_damage.jpg", "Hallway water damage - Okafor claim")
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


def build_property_incomplete(root: Path) -> None:
    # Missing mandatory: inspection_report, receipts, policy_declaration.
    _doc(
        root / "fnol" / "notice_of_loss.docx",
        "Notice of Loss - Property",
        [
            "Policy Number: PROP-8820-1156",
            "Carrier Name: Harborstone Insurance Group",
            "Date of Loss: 06/10/2026",
            "Cause of Loss: Wind damage during severe thunderstorm.",
            "Loss Description: High winds tore several shingles from the roof and "
            "cracked a rear window; minor water intrusion into the attic.",
            "Property Address: 902 Larkspur Ct, Millbrook, GA",
            "Property Type: Single-family residence",
            "Occupancy Status: Owner-occupied, primary residence",
        ],
    )
    _doc(
        root / "inventory_list" / "inventory_list.docx",
        "Damaged Contents Inventory",
        [
            "Inventory Reference: INV-2026-5561",
            "Damaged Items List: Roof shingles (approx. 30 sq ft affected), rear "
            "bedroom window, attic insulation (partial water staining).",
        ],
    )
    _photo(root / "evidence" / "roof_shingle_damage.jpg", "Roof shingle damage - Millbrook claim")
    _irrelevant_unsupported(root / "irrelevant" / "unrelated_expense_log.csv")


def build_property_conflicting_values(root: Path) -> None:
    _doc(
        root / "fnol" / "notice_of_loss.docx",
        "Notice of Loss - Property",
        [
            "Policy Number: PROP-3315-7729",
            "Carrier Name: Harborstone Insurance Group",
            "Policy Effective Dates: 09/15/2025 to 09/15/2026",
            "Date of Loss: 01/28/2026",
            "Cause of Loss: Burst pipe in second-floor bathroom during a hard freeze.",
            "Loss Description: Water traveled through the second-floor bathroom wall "
            "into the kitchen ceiling below, causing drywall and cabinetry damage on "
            "both floors.",
            "Property Address: 47 Winslow Terrace, Hartford, CT",
            "Property Type: Single-family residence",
            "Occupancy Status: Owner-occupied, primary residence",
        ],
    )
    _doc(
        root / "inspection_report" / "inspection_report.docx",
        "Harborstone Field Inspection Report",
        [
            "Inspection Report Reference: HFI-2026-2207",
            "Findings: Active leak traced to a burst supply line behind the "
            "second-floor bathroom vanity. Water damage extends through the subfloor "
            "into the kitchen ceiling directly below.",
        ],
    )
    _doc(
        root / "inventory_list" / "inventory_list.docx",
        "Damaged Contents Inventory",
        [
            "Inventory Reference: INV-2026-2207-A",
            "Damaged Items List: Bathroom vanity and subfloor, kitchen ceiling drywall, "
            "kitchen upper cabinets (water staining), hallway ceiling.",
        ],
    )
    _doc(
        root / "receipts" / "receipts_initial_adjuster_estimate.docx",
        "Initial Adjuster Loss Estimate",
        [
            "Probable Amount of Loss: $22,400.00",
            "Replacement Cost Estimate: $24,900.00",
        ],
    )
    _doc(
        root / "receipts" / "receipts_contractor_revised_estimate.docx",
        "Contractor Revised Estimate (Supplemental)",
        [
            "Probable Amount of Loss: $31,150.00",
            "Replacement Cost Estimate: $33,600.00",
            "Note: Revised upward after opening the wall revealed additional mold "
            "remediation was required behind the kitchen cabinetry.",
        ],
    )
    _doc(
        root / "policy_declaration" / "policy_declaration.docx",
        "Harborstone Insurance Group - Policy Declarations",
        [
            "Policy Number: PROP-3315-7729",
            "Named Insured: Sofia Marchetti",
            "Policy Effective Dates: 09/15/2025 to 09/15/2026",
            "Coverage: Dwelling, Personal Property, Water Damage Rider",
        ],
    )
    _photo(root / "evidence" / "kitchen_ceiling_water_damage.jpg", "Kitchen ceiling damage - Marchetti claim")
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


# ---------------------------------------------------------------------------
# HEALTH
# ---------------------------------------------------------------------------


def build_health_complete_clean(root: Path) -> None:
    _doc(
        root / "pre_authorization" / "pre_authorization.docx",
        "Pre-Authorization Approval",
        [
            "Pre-Authorization Number: PA-2026-778341",
            "Patient Name: Julia Ferns",
            "Member ID: HLX-2291-8804",
            "Date of Birth: 08/14/1989",
            "Relationship to Subscriber: Self",
            "Provider Name: Lakeshore Orthopedic Associates",
            "Provider NPI: 1447382910",
            "Network Status: In-network",
            "Service Requested: Right knee arthroscopic meniscus repair.",
        ],
    )
    _doc(
        root / "discharge_summary" / "discharge_summary.docx",
        "Discharge Summary",
        [
            "Discharge Summary Reference: DS-2026-778341",
            "Patient Name: Julia Ferns",
            "Admission Date: 05/03/2026",
            "Discharge Date: 05/03/2026",
            "Date of Service: 05/03/2026",
            "ICD-10 Codes: M23.221",
            "CPT Codes: 29882",
            "Diagnosis Description: Tear of medial meniscus, right knee, current injury.",
            "Procedure: Outpatient arthroscopic partial meniscectomy, right knee. "
            "Patient tolerated the procedure well and was discharged same day in "
            "stable condition with standard post-op instructions.",
        ],
    )
    _doc(
        root / "itemized_bill" / "itemized_bill.docx",
        "Itemized Billing Statement",
        [
            "Provider Name: Lakeshore Orthopedic Associates",
            "Patient Name: Julia Ferns",
            "Billed Amount: $14,820.00",
            "Co-Pay Amount: $250.00",
            "Allowed Amount: $9,640.00",
            "Line items: Facility fee, surgeon fee, anesthesia, post-op supplies.",
        ],
    )
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


def build_health_incomplete(root: Path) -> None:
    # Missing mandatory: discharge_summary, itemized_bill.
    _doc(
        root / "pre_authorization" / "pre_authorization.docx",
        "Pre-Authorization Approval",
        [
            "Pre-Authorization Number: PA-2026-551209",
            "Patient Name: Marcus Delgado",
            "Member ID: HLX-6634-2217",
            "Date of Birth: 11/02/1975",
            "Relationship to Subscriber: Self",
            "Provider Name: Cedar Valley Cardiology",
            "Provider NPI: 1932847710",
            "Network Status: In-network",
            "Service Requested: Cardiac stress test with echocardiogram.",
        ],
    )
    _irrelevant_unsupported(root / "irrelevant" / "unrelated_expense_log.csv")
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


def build_health_conflicting_values(root: Path) -> None:
    _doc(
        root / "pre_authorization" / "pre_authorization.docx",
        "Pre-Authorization Approval",
        [
            "Pre-Authorization Number: PA-2026-902247",
            "Patient Name: Renee Ashworth",
            "Member ID: HLX-1180-4423",
            "Date of Birth: 02/27/1993",
            "Relationship to Subscriber: Spouse",
            "Provider Name: Northgate Surgical Center",
            "Provider NPI: 1558820047",
            "Network Status: In-network",
            "Service Requested: Laparoscopic appendectomy.",
        ],
    )
    _doc(
        root / "discharge_summary" / "discharge_summary.docx",
        "Discharge Summary",
        [
            "Discharge Summary Reference: DS-2026-902247",
            "Patient Name: Renee Ashworth",
            "Admission Date: 07/11/2026",
            "Discharge Date: 07/12/2026",
            "Date of Service: 07/11/2026",
            "ICD-10 Codes: K35.80",
            "CPT Codes: 44970",
            "Diagnosis Description: Unspecified acute appendicitis.",
            "Procedure: Laparoscopic appendectomy, uncomplicated. Overnight "
            "observation for post-operative monitoring, discharged the following morning.",
        ],
    )
    _doc(
        root / "itemized_bill" / "itemized_bill_facility.docx",
        "Itemized Billing Statement - Facility",
        [
            "Provider Name: Northgate Surgical Center",
            "Patient Name: Renee Ashworth",
            "Billed Amount: $19,340.00",
            "Co-Pay Amount: $150.00",
        ],
    )
    _doc(
        root / "itemized_bill" / "itemized_bill_corrected.docx",
        "Itemized Billing Statement - Corrected",
        [
            "Provider Name: Northgate Surgical Center",
            "Patient Name: Renee Ashworth",
            "Billed Amount: $22,875.00",
            "Co-Pay Amount: $150.00",
            "Note: Corrected statement -- original omitted overnight observation "
            "facility charges.",
        ],
    )
    _irrelevant_docx(root / "irrelevant" / "office_allhands_notes.docx")


# ---------------------------------------------------------------------------


BUILDERS = {
    "auto_complete_clean": build_auto_complete_clean,
    "auto_incomplete": build_auto_incomplete,
    "auto_conflicting_values": build_auto_conflicting_values,
    "property_complete_clean": build_property_complete_clean,
    "property_incomplete": build_property_incomplete,
    "property_conflicting_values": build_property_conflicting_values,
    "health_complete_clean": build_health_complete_clean,
    "health_incomplete": build_health_incomplete,
    "health_conflicting_values": build_health_conflicting_values,
}


def main() -> None:
    if OUT_ROOT.exists():
        shutil.rmtree(OUT_ROOT)
    OUT_ROOT.mkdir(parents=True)

    for name, builder in BUILDERS.items():
        folder = OUT_ROOT / name
        builder(folder)
        print(f"built {folder.relative_to(REPO_ROOT)}")

    readme = OUT_ROOT / "README.md"
    readme.write_text(
        "# Sample claim folders\n\n"
        "Nine synthetic claim folders (3 per LOB) for exercising the full "
        "pipeline end-to-end, generated by `scripts/generate_sample_claims.py`. "
        "Upload any one of these folders whole (via the \"Browse folder\" button "
        "on the Start Claim page) to test folder upload with nested subdirectories.\n\n"
        "| Folder | LOB | Scenario |\n"
        "|---|---|---|\n"
        "| `auto_complete_clean/` | Auto | All mandatory docs present, fields agree -- expect STP candidate |\n"
        "| `auto_incomplete/` | Auto | Missing police report + policy declaration -- expect needs_review/high_risk |\n"
        "| `auto_conflicting_values/` | Auto | Two repair estimates disagree on dollar amount; includes a 22-page supplemental PDF |\n"
        "| `property_complete_clean/` | Property | All mandatory docs present, fields agree -- expect STP candidate |\n"
        "| `property_incomplete/` | Property | Missing inspection report, receipts, policy declaration |\n"
        "| `property_conflicting_values/` | Property | Initial vs. revised contractor estimate disagree on loss amount |\n"
        "| `health_complete_clean/` | Health | All mandatory docs present, fields agree -- expect STP candidate |\n"
        "| `health_incomplete/` | Health | Missing discharge summary + itemized bill |\n"
        "| `health_conflicting_values/` | Health | Original vs. corrected itemized bill disagree on billed amount |\n\n"
        "Every folder also includes at least one irrelevant file under `irrelevant/` "
        "-- either off-topic but readable (an all-hands notes DOCX) or an unsupported "
        "extension (a stray .csv) -- to confirm the pipeline classifies/ignores noise "
        "rather than getting confused by it.\n",
    )
    print(f"wrote {readme.relative_to(REPO_ROOT)}")


if __name__ == "__main__":
    main()
